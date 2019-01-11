#This is a basic version of a script that will merge the contents of any word documents given. The user will select a current word
#document and then the user decides if they would like to copy the contents into a new word document, or to append it to the end of an
#existing document. Currently this script only copys the text and not the style, in the future however I will be updating that as well as
#adding GUI.

from docx import Document

doc = input("What is the document that you would like to copy from?: ")
docname = (doc+".docx")
currentdoc = Document(docname)

content = []
for para in currentdoc.paragraphs:
    c = para.text
    content.append(c)

nameofnewdoc = input("\nWould you like to save to a new word document: ")
if nameofnewdoc.title() == "Yes":
    name = input("\nWhat would you like to name the new document?: ")
    output = Document()
    for item in content:
        output.add_paragraph(item)
    output.save((name+".docx"))
else:
    existing = input("\nWhat document would you like to save it onto: ")
    output = Document(existing+".docx")
    for item in content:
        output.add_paragraph(item)
    output.save((existing+".docx"))
    
    

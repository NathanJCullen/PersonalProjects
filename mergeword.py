from docx import Document

doc = input("What is the document that you would like to copy from?: ")
docname = (doc+".docx")
currentdoc = Document(docname)

content = []
for para in currentdoc.paragraphs:
    c = para.text
    content.append(c)

nameofnewdoc = input("\nWould you like to save to a new word document: ")
if nameofnewdoc.title() == "Yes":
    name = input("\nWhat would you like to name the new document?: ")
    output = Document()
    for item in content:
        output.add_paragraph(item)
    output.save((name+".docx"))
else:
    existing = input("\nWhat document would you like to save it onto: ")
    output = Document(existing+".docx")
    for item in content:
        output.add_paragraph(item)
    output.save((existing+".docx"))
    
    
